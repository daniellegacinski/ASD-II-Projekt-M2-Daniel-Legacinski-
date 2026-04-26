from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import csv, glob

BASE = Path(__file__).resolve().parents[1]
DOCS = BASE / "docs"
OUTS = sorted((BASE / "outputs").glob("experiments_cli_*"))
EXP = OUTS[-1] if OUTS else None


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)


def add_fig(doc, path, caption, width=5.4):
    if not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(135, 37, 48)
    return h


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('Projekt M-II\nChaotyczne przeksztalcanie obrazu cyfrowego')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(135,37,48)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Dokumentacja projektu\n').bold = True
    meta.add_run('Opracowal: Daniel Legacinski\n')
    meta.add_run('Technologia: Python, NumPy, Pillow, Matplotlib, Tkinter\n')
    meta.add_run('Rok: 2026')
    doc.add_page_break()

    add_heading(doc, '1. Cel projektu', 1)
    doc.add_paragraph('Celem projektu jest praktyczne pokazanie mechanizmow przeksztalcania obrazu cyfrowego. Projekt nie ma tworzyc bezpiecznego szyfru, tylko eksperyment dydaktyczny pokazujacy roznice pomiedzy prosta permutacja, substytucja, pseudo-losowoscia i wizualnym chaosem.')
    doc.add_paragraph('W projekcie zostaly zrealizowane trzy obowiazkowe etapy. Kazdy etap ma osobny algorytm scramblingu oraz osobny algorytm odwrotny, czyli unscrambling. Odwracanie korzysta tylko z obrazu wynikowego oraz klucza/parametrow. Program nie zapisuje dodatkowej mapy pomocniczej.')
    doc.add_paragraph('Najwazniejsze zalozenie: jezeli klucz i parametry sa poprawne, obraz musi zostac odtworzony piksel po pikselu. Jezeli klucz jest bledny, program powinien pokazac, ze wynik odtworzenia jest niepoprawny.')

    add_heading(doc, '2. Struktura projektu', 1)
    items = [
        ('run_gui.py', 'glowna aplikacja graficzna GUI'),
        ('run_experiments.py', 'uruchamianie eksperymentow bez GUI i zapis wynikow'),
        ('test_reversibility.py', 'test poprawnej odwracalnosci dla etapow 1, 2 i 3'),
        ('scrambler/algorithms.py', 'implementacja algorytmow scramblingu i unscramblingu'),
        ('scrambler/metrics.py', 'metryki: korelacja sasiadow, MSE, MAE, PSNR, roznice pikseli'),
        ('assets/samples', 'obrazy testowe: obraz naturalny i obraz o silnej strukturze'),
        ('outputs', 'wyniki eksperymentow, obrazy, CSV i wykresy')]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Plik/folder'
    table.rows[0].cells[1].text = 'Znaczenie'
    for cell in table.rows[0].cells: set_cell_shading(cell, 'E8D7DA')
    for a,b in items:
        row = table.add_row().cells; row[0].text = a; row[1].text = b

    add_heading(doc, '3. Reprezentacja obrazu cyfrowego', 1)
    doc.add_paragraph('Obraz jest traktowany jako macierz pikseli RGB. Kazdy piksel ma trzy skladowe: R, G i B, a kazda skladowa jest liczba calkowita z zakresu 0-255. Po wczytaniu obraz jest konwertowany do formatu RGB, aby wszystkie algorytmy pracowaly na tej samej strukturze danych.')
    doc.add_paragraph('Dla obrazu o wysokosci H i szerokosci W liczba pikseli wynosi N = H * W. W Etapie 2 i 3 obraz jest tymczasowo zamieniany na jednowymiarowa tablice pikseli, aby mozna bylo wykonac permutacje indeksow.')
    add_code(doc, 'flat = arr.reshape(H * W, 3)\n# jeden element flat[i] oznacza jeden piksel RGB')

    add_heading(doc, '4. Etap 1 - naiwny scrambling', 1)
    doc.add_paragraph('Etap 1 jest celowo prosty. Program wykonuje deterministyczne przesuniecia wierszy oraz kolumn. Parametr liczbowy okresla krok przesuniecia, a klucz jest zamieniany na seed. Ten etap jest odwracalny, ale nie usuwa dobrze struktury obrazu.')
    doc.add_paragraph('Scrambling: dla kazdego wiersza wykonywane jest przesuniecie cykliczne w prawo, a dla kazdej kolumny przesuniecie cykliczne w dol. Unscrambling wykonuje te same operacje w odwrotnej kolejnosci i z przeciwnym znakiem przesuniecia.')
    add_code(doc, 'for r in range(H):\n    row[r] = roll(row[r], shift_r)\nfor c in range(W):\n    column[c] = roll(column[c], shift_c)\n# odwrotnie: najpierw kolumny w gore, potem wiersze w lewo')
    doc.add_paragraph('Slabosc metody: tekst, szachownica, gradienty i krawedzie moga nadal byc czesciowo widoczne, poniewaz piksele w lokalnych strukturach sa tylko przesuwane, a nie mieszane globalnie.')

    add_heading(doc, '5. Etap 2 - czysta permutacja sterowana kluczem', 1)
    doc.add_paragraph('Etap 2 realizuje czysta permutacje pikseli. Wartosc piksela nie jest zmieniana, zmienia sie tylko jego pozycja. Permutacja jest generowana algorytmem Fishera-Yatesa sterowanym jawnym generatorem LCG. Dzięki temu dla tego samego klucza i parametru zawsze powstaje ta sama permutacja.')
    doc.add_paragraph('Formalnie permutacja jest funkcja P: {0,...,N-1} -> {0,...,N-1}. Algorytm odwrotny korzysta z tej samej permutacji i rozklada piksele z powrotem na poprzednie indeksy.')
    add_code(doc, 'scrambled[i] = original[P[i]]\nrestored[P[i]] = scrambled[i]\n# dlatego P^(-1)(P(i)) = i')
    doc.add_paragraph('Czysta permutacja zmniejsza widocznosc lokalnych struktur, ale sama nie jest bezpieczna. Histogram kolorow pozostaje identyczny, bo wartosci pikseli nie sa zmieniane.')

    add_heading(doc, '6. Etap 3 - mechanizm wzmacniajacy', 1)
    doc.add_paragraph('W Etapie 3 zastosowano klase C: hybryda, czyli permutacja + substytucja. Najpierw obraz jest permutowany jak w Etapie 2, a potem do kazdej skladowej RGB dodawany jest strumien wartosci 0-255 modulo 256. Strumien jest generowany deterministycznie z klucza, parametru i numeru rundy.')
    doc.add_paragraph('Operacja substytucji jest odwracalna, poniewaz dodawanie modulo 256 mozna cofnac odejmowaniem modulo 256. Przy odtwarzaniu trzeba wykonac operacje w odwrotnej kolejnosci: najpierw odjac strumien, potem cofnac permutacje.')
    add_code(doc, 'scramble:   y = (P(x) + stream) mod 256\nunscramble: x = P^(-1)((y - stream) mod 256)')
    doc.add_paragraph('Mechanizm poprawia efekt wizualny, bo zmienia nie tylko polozenie pikseli, ale tez ich wartosci. Nadal nie jest to bezpieczny szyfr, poniewaz generator LCG i funkcja seed sa proste oraz mozliwe do analizy.')

    add_heading(doc, '7. Interfejs uzytkownika GUI', 1)
    doc.add_paragraph('Aplikacja GUI zostala wykonana w Tkinter. Interfejs pozwala wczytac obraz PNG, JPEG lub BMP, wybrac etap 1/2/3, wpisac klucz, podac parametr i liczbe rund dla Etapu 3. Program jednoczesnie pokazuje obraz oryginalny, obraz przeksztalcony oraz obraz odtworzony.')
    doc.add_paragraph('Dostepne przyciski: Wczytaj obraz, Scramble, Unscramble - poprawny klucz, Unscramble - bledny klucz, Zapisz aktualne wyniki oraz Uruchom pelne eksperymenty. GUI jest narzedziem analitycznym, a nie tylko dekoracja.')

    add_heading(doc, '8. Metryki eksperymentalne', 1)
    doc.add_paragraph('Projekt liczy dwie obowiazkowe grupy metryk. Pierwsza to korelacja sasiadujacych pikseli przed i po scramblingu. Druga to roznica obrazu po unscramblingu blednym kluczem. Dodatkowo program liczy MSE, MAE, maksymalna roznice oraz PSNR.')
    metrics = [('Korelacja sasiadow', 'pokazuje, czy piksele obok siebie sa nadal podobne'), ('MSE', 'sredni kwadrat roznicy miedzy obrazami'), ('MAE', 'srednia bezwzgledna roznica pikseli'), ('Max diff', 'najwieksza roznica dla pojedynczej skladowej'), ('PSNR', 'miara podobienstwa obrazow; nieskonczona dla identycznych obrazow')]
    t = doc.add_table(rows=1, cols=2); t.style='Table Grid'
    t.rows[0].cells[0].text='Metryka'; t.rows[0].cells[1].text='Interpretacja'
    for cell in t.rows[0].cells: set_cell_shading(cell,'E8D7DA')
    for a,b in metrics:
        row=t.add_row().cells; row[0].text=a; row[1].text=b

    add_heading(doc, '9. Wyniki eksperymentow', 1)
    if EXP:
        add_fig(doc, BASE / 'assets/samples/sample_structured.png', 'Rys. 1. Obraz testowy o silnej strukturze.', width=4.8)
        add_fig(doc, EXP / 'sample_structured_stage1_scrambled.png', 'Rys. 2. Etap 1 - widoczne pozostalosci struktury po naiwnym przesunieciu.', width=4.8)
        add_fig(doc, EXP / 'sample_structured_stage2_scrambled.png', 'Rys. 3. Etap 2 - czysta permutacja pikseli.', width=4.8)
        add_fig(doc, EXP / 'sample_structured_stage3_scrambled.png', 'Rys. 4. Etap 3 - permutacja + substytucja modularna.', width=4.8)
        add_fig(doc, EXP / 'sample_structured_stage3_restored_ok.png', 'Rys. 5. Poprawne odtworzenie obrazu dla Etapu 3.', width=4.8)
        add_fig(doc, EXP / 'sample_structured_stage3_restored_wrong_key.png', 'Rys. 6. Proba odtworzenia Etapu 3 blednym kluczem.', width=4.8)
        add_fig(doc, EXP / 'correlation_chart.png', 'Rys. 7. Porownanie korelacji po scramblingu.', width=6.0)
        # tabela wynikow
        csv_path = EXP / 'metrics.csv'
        if csv_path.exists():
            with open(csv_path, encoding='utf-8') as f:
                rows = list(csv.DictReader(f))
            doc.add_paragraph('Tabela ponizej zawiera skrot wynikow z automatycznego eksperymentu. Pelny plik CSV znajduje sie w folderze outputs.')
            for r in rows:
                line = (f"Obraz: {r['image']}, etap {r['stage']} | korelacja po H: "
                        f"{float(r['corr_scrambled_horizontal']):.4f} | odtworzenie OK: {r['restore_identical']} | "
                        f"MSE bledny klucz: {float(r['wrong_key_mse']):.2f} | MAE bledny klucz: {float(r['wrong_key_mae']):.2f}")
                doc.add_paragraph(line, style=None)
        doc.add_page_break()

    add_heading(doc, '10. Analiza porownawcza etapow', 1)
    doc.add_paragraph('Etap 1 jest najprostszy i pokazuje porazke metody naiwnej. Obraz wyglada inaczej, ale struktura moze nadal przetrwac. Przy obrazie z tekstem lub szachownica mozna rozpoznac regularnosc, poniewaz algorytm przesuwa cale wiersze i kolumny.')
    doc.add_paragraph('Etap 2 usuwa wiele lokalnych struktur, poniewaz piksele trafiaja w dalekie miejsca. Jednak wartosci pikseli pozostaja takie same, dlatego histogram obrazu i ogolna paleta kolorow sie nie zmieniaja. To oznacza, ze sama permutacja nie zapewnia bezpieczenstwa.')
    doc.add_paragraph('Etap 3 daje najsilniejszy efekt wizualny, bo laczy zmiane pozycji i zmiane wartosci pikseli. Bledny klucz powoduje duze roznice w obrazie odtworzonym. Mimo tego nie mozna twierdzic, ze jest to szyfr bezpieczny, bo uzyty generator jest prosty i przewidywalny.')

    add_heading(doc, '11. Krytyka wlasnego rozwiazania: dlaczego to NIE jest bezpieczny szyfr', 1)
    doc.add_paragraph('1. Jakie informacje atakujacy nadal moze odzyskac? W Etapie 1 atakujacy moze zobaczyc resztki struktury, bo przesuniecia zachowuja lokalne zaleznosci. W Etapie 2 atakujacy widzi histogram kolorow, poniewaz wartosci pikseli sie nie zmieniaja. W Etapie 3 histogram jest zmieniony, ale generator LCG jest prosty i mozliwy do odtworzenia przy analizie wielu danych.')
    doc.add_paragraph('2. Co dzieje sie przy ataku typu known-plaintext? Jezeli atakujacy zna obraz oryginalny i wynikowy dla tego samego klucza, moze probowac odzyskac permutacje albo parametry generatora. W Etapie 2 jest to szczegolnie grozne, bo permutacja bezposrednio laczy piksele oryginalne z wynikowymi. W Etapie 3 atak jest trudniejszy, ale nadal mozliwy, bo substytucja jest prosta i deterministyczna.')
    doc.add_paragraph('3. Dlaczego zwiekszanie chaosu nie rozwiazuje problemu? Obraz moze wygladac losowo, ale wyglad losowy nie oznacza bezpieczenstwa. Bezpieczenstwo wymaga odpornego projektu kryptograficznego, analizy klucza, odpornego generatora i testow kryptograficznych. Ten projekt celowo tego nie zapewnia, poniewaz ma charakter dydaktyczny.')

    add_heading(doc, '12. Instrukcja uruchomienia', 1)
    doc.add_paragraph('Krok 1. Rozpakuj folder projektu. Krok 2. Otworz terminal w tym folderze. Krok 3. Zainstaluj biblioteki poleceniem pip install -r requirements.txt. Krok 4. Uruchom aplikacje poleceniem python run_gui.py.')
    add_code(doc, 'pip install -r requirements.txt\npython run_gui.py')
    doc.add_paragraph('Aby wykonac eksperymenty automatycznie bez GUI, uruchom:')
    add_code(doc, 'python run_experiments.py')
    doc.add_paragraph('Aby sprawdzic odwracalnosc wszystkich etapow, uruchom:')
    add_code(doc, 'python test_reversibility.py')

    add_heading(doc, '13. Zalacznik techniczny - pseudokod algorytmow', 1)
    doc.add_paragraph('Pseudokod ponizej pokazuje, ze kazdy etap ma osobna procedure odwrotna. Jest to wazne, poniewaz wymaganie projektu zabrania przechowywania dodatkowych danych pomocniczych. Program musi odtworzyc obraz tylko z obrazu wynikowego oraz klucza/parametrow.')
    add_code(doc, '''ETAP 1 SCRAMBLE:
  dla kazdego wiersza r: przesun wiersz o shift_r
  dla kazdej kolumny c: przesun kolumne o shift_c
ETAP 1 UNSCRAMBLE:
  dla kazdej kolumny c od konca: przesun kolumne o -shift_c
  dla kazdego wiersza r od konca: przesun wiersz o -shift_r''')
    add_code(doc, '''ETAP 2 SCRAMBLE:
  seed = funkcja(klucz, parametr)
  P = FisherYates(N, seed)
  wynik[i] = obraz[P[i]]
ETAP 2 UNSCRAMBLE:
  P = FisherYates(N, seed)
  odtworzony[P[i]] = wynik[i]''')
    add_code(doc, '''ETAP 3 SCRAMBLE:
  dla kazdej rundy:
    obraz = permutacja(obraz)
    obraz = (obraz + strumien) mod 256
ETAP 3 UNSCRAMBLE:
  dla kazdej rundy w odwrotnej kolejnosci:
    obraz = (obraz - strumien) mod 256
    obraz = odwrotna_permutacja(obraz)''')
    doc.add_paragraph('Wazna obserwacja: odwracalnosc nie oznacza bezpieczenstwa. Odwracalnosc jest tylko warunkiem technicznym, aby poprawny uzytkownik mogl odzyskac obraz. Bezpieczenstwo wymagaloby odpornosci na analize atakujacego, czego projekt celowo nie zapewnia.')
    doc.add_page_break()

    add_heading(doc, '14. Wnioski', 1)
    doc.add_paragraph('Projekt pokazuje, ze przeksztalcanie obrazu moze byc w pelni odwracalne bez przechowywania danych pomocniczych. Najwazniejszy warunek to jednoznaczny klucz, parametr i algorytm odwrotny. Etap 1 pokazuje slabosc metod naiwnych, Etap 2 pokazuje ograniczenia samej permutacji, a Etap 3 pokazuje, ze dodanie substytucji wzmacnia efekt wizualny.')
    doc.add_paragraph('Najwazniejszy wniosek jest taki, ze wizualny chaos nie jest tym samym co bezpieczenstwo. Nawet jezeli obraz wyglada jak przypadkowy szum, algorytm moze byc latwy do przeanalizowania. Dlatego projekt nalezy traktowac jako eksperyment edukacyjny, a nie jako narzedzie do ochrony danych.')

    add_heading(doc, '15. Bibliografia i zrodla', 1)
    for src in [
        'Dokumentacja Python: https://docs.python.org/3/',
        'Dokumentacja NumPy: https://numpy.org/doc/',
        'Dokumentacja Pillow: https://pillow.readthedocs.io/',
        'Dokumentacja Matplotlib: https://matplotlib.org/stable/',
        'Opis permutacji Fisher-Yates: https://en.wikipedia.org/wiki/Fisher%E2%80%93Yates_shuffle',
        'Opis liniowego generatora kongruencyjnego LCG: https://en.wikipedia.org/wiki/Linear_congruential_generator'
    ]:
        doc.add_paragraph(src, style=None)

    path = DOCS / 'Dokumentacja_Projekt_M-II_Daniel_Legacinski.docx'
    doc.save(path)
    print(path)

if __name__ == '__main__':
    main()
